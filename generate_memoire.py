# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0, 0, 0)

doc = Document()

# ── marges ────────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

def add_page_numbers(doc):
    """Section 0 = page de titre (pas de numero).
       Sections 1+ = numerotation a partir de 1."""
    sections = doc.sections
    for i, section in enumerate(sections):
        if i == 0:
            # Page de titre : pied de page vide
            footer = section.footer
            footer.is_linked_to_previous = False
            if footer.paragraphs:
                footer.paragraphs[0].clear()
        else:
            # A partir du Resume : numerotation depuis 1
            add_page_number_to_section(section, start=1)

# ── styles de base ────────────────────────────────────────────────────────────
def set_font(run, size=12, bold=False, italic=False):
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.color.rgb = BLACK

def heading(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    sizes = {1: 16, 2: 14, 3: 13, 4: 12}
    set_font(run, size=sizes.get(level, 12), bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    return p

def para(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(1.0)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, size=12)
    p.paragraph_format.space_after = Pt(3)
    return p

def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=11, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=11)
    doc.add_paragraph()

def page_break():
    doc.add_page_break()

def section_break_after_title():
    """Cree un saut de section apres la page de titre."""
    p = doc.add_paragraph()
    pPr = OxmlElement('w:pPr')
    sectPr = OxmlElement('w:sectPr')
    # Type: nextPage
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), '11906')
    pgSz.set(qn('w:h'), '16838')
    sectPr.append(pgSz)
    pPr.append(sectPr)
    p._p.append(pPr)

def add_page_number_to_section(section, start=1):
    """Ajoute un numero de page centre dans le pied de page d'une section."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK

    # Numero de page debut
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    # Demarrer la numerotation a partir de 'start'
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:start'), str(start))
    section._sectPr.append(pgNumType)

BASE = 'c:/Users/tassili/Desktop/pfe/'

def image(path, caption=None, width_cm=15):
    """Insere une image centree avec une legende optionnelle."""
    import os
    if not os.path.exists(path):
        p = doc.add_paragraph()
        run = p.add_run(f"[Image manquante : {path}]")
        set_font(run, size=10, italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption)
        set_font(cr, size=10, italic=True)
        cp.paragraph_format.space_after = Pt(10)

def center(text, size=12, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(6)
    return p

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE DE TITRE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
center("REPUBLIQUE ALGERIENNE DEMOCRATIQUE ET POPULAIRE", size=11)
center("Ministere de l'Enseignement Superieur et de la Recherche Scientifique", size=11)
doc.add_paragraph()
center("Universite — Departement d'Informatique", size=12, bold=True)
doc.add_paragraph()
doc.add_paragraph()
center("Projet de fin d'etudes pour l'obtention du diplome LICENCE", size=14, bold=True)
center("Universite de Sciences et Technologies Houari Boumediene — USTHB", size=13, bold=True)
center("Faculte d'Informatique", size=12)
center("Specialite : Informatique Generale", size=12)
doc.add_paragraph()
doc.add_paragraph()
center("CONCEPTION ET DEVELOPPEMENT D'UNE APPLICATION MOBILE", size=14, bold=True)
center("DE RESERVATION DES BILLETS D'AVION", size=14, bold=True)
doc.add_paragraph()
center("Backend Django  |  Base de Donnees MySQL  |  API Duffel", size=11)
doc.add_paragraph()
doc.add_paragraph()
center("Presente par :", size=12)
center("LADRAA MOHAMED BADREDDINE", size=12, bold=True)
center("BENFATTOUM MOHAMED", size=12, bold=True)
center("KEZOUIT OMAR", size=12, bold=True)
center("ZAROUTA MOHAMED RACIM", size=12, bold=True)
doc.add_paragraph()
center("Encadre par :", size=12)
center("[Nom Encadrant]", size=12, bold=True)
doc.add_paragraph()
doc.add_paragraph()
center("Annee universitaire : 2024 / 2025", size=12, bold=True)

section_break_after_title()

# ══════════════════════════════════════════════════════════════════════════════
#  RESUME
# ══════════════════════════════════════════════════════════════════════════════
heading("RESUME", 1)
para("Ce memoire presente le developpement d'une application mobile de reservation de vols, realisee dans le cadre du projet de fin d'etudes en Licence Informatique. L'application repose sur une architecture client/serveur moderne combinant Flutter pour le frontend mobile, Django REST Framework pour le backend, MySQL comme systeme de gestion de base de donnees, et l'API Duffel pour l'acces en temps reel aux offres de vols des compagnies aeriennes mondiales.")
para("Le systeme permet a trois types d'utilisateurs — client, agent et administrateur — d'interagir avec la plateforme selon leurs roles respectifs. Le client recherche et reserve des vols, l'agent confirme les reservations et gere les demandes en attente, et l'administrateur supervise l'ensemble du systeme, gere les utilisateurs et configure les codes promotionnels.")
para("Mots-cles : Application mobile, Flutter, Django, API REST, Duffel API, MySQL, reservation de vols, architecture client/serveur.", italic=True)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  TABLE DES MATIERES
# ══════════════════════════════════════════════════════════════════════════════
heading("TABLE DES MATIERES", 1)
toc_items = [
    "Introduction Generale",
    "Chapitre 1 : Generalites",
    "    I. Introduction",
    "    II. Generalites sur le Transport Aerien",
    "    III. Applications Mobiles",
    "    IV. Architecture Client/Serveur",
    "    V. UML",
    "Chapitre 2 : Analyse et Conception",
    "    I. Introduction",
    "    II. Analyse des Besoins",
    "    III. Diagramme de Cas d'Utilisation",
    "    IV. Diagramme de Classes",
    "    V. Diagrammes de Sequence",
    "    VI. Schema de la Base de Donnees",
    "Chapitre 3 : Realisation",
    "    I. Introduction",
    "    II. Environnement de Developpement",
    "    III. Architecture de l'Application",
    "    IV. Fonctionnement des Modules",
    "    V. Algorithme A priori — Systeme de Recommandation",
    "    VI. Tests et Validation",
    "Conclusion Generale",
    "References Bibliographiques",
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    set_font(run, size=12)
    p.paragraph_format.space_after = Pt(3)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  INTRODUCTION GENERALE
# ══════════════════════════════════════════════════════════════════════════════
heading("INTRODUCTION GENERALE", 1)
para("Au cours des dernieres decennies, l'informatique a transforme le secteur des services par sa capacite exceptionnelle a gerer des volumes de donnees complexes. Cette influence s'etend desormais a de nombreux domaines, dont celui du transport aerien, qui constitue le coeur de notre projet de fin d'etudes.")
para("La notion de reservation est centrale dans le secteur de l'aviation. Les compagnies aeriennes proposent et gerent leurs vols avant de les proposer aux passagers. Gerer ces reservations est au coeur des preoccupations des gestionnaires pour plusieurs raisons. L'accessibilite aux tarifs en temps reel et la gestion des flux de passagers coutent cher et demandent une grande precision. Le secteur a donc tout interet a optimiser son systeme de reservation.")
para("Nous nous interessons plus precisement a la reservation des billets d'avion, afin d'ameliorer les differents traitements effectues pour les voyageurs. La gestion des reservations de la plupart des agences traditionnelles s'effectue parfois de facon complexe ou peu accessible, engendrant des problemes tels que la lenteur dans l'acces aux tarifs et le manque de flexibilite pour l'utilisateur.")
para("Pour faciliter le travail administratif et l'experience utilisateur, nous proposons dans le cadre de ce projet la conception et la realisation d'une application mobile, permettant aux passagers et aux administrateurs de gerer les vols et les reservations via l'API Duffel. Cette solution repose sur les technologies modernes : Flutter pour l'interface mobile, Django REST Framework pour le serveur, MySQL pour la persistance des donnees, et l'algorithme A priori pour les recommandations intelligentes de destinations.")
para("Ce memoire s'organise en trois chapitres principaux. Le premier chapitre pose les bases conceptuelles et technologiques du projet. Le deuxieme chapitre presente l'analyse des besoins et la conception UML du systeme. Le troisieme chapitre decrit la realisation concrete de l'application, incluant l'integration de l'algorithme de fouille de donnees A priori.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPITRE 1
# ══════════════════════════════════════════════════════════════════════════════
heading("CHAPITRE 1 : GENERALITES", 1)

heading("I. Introduction", 2)
para("Le premier chapitre a pour objectif de presenter le contexte general dans lequel s'inscrit notre projet. Nous aborderons successivement les fondements du transport aerien moderne, les specificites des applications mobiles, l'architecture client/serveur qui structure notre systeme, et les concepts UML utilises pour modeliser notre application.")

heading("II. Generalites sur le Transport Aerien", 2)
heading("II.1. Evolution du secteur aerien", 3)
para("Le transport aerien est l'un des secteurs economiques qui a connu la croissance la plus spectaculaire du XXe siecle. Depuis les premiers vols commerciaux dans les annees 1920, l'aviation civile n'a cesse de se developper pour devenir aujourd'hui un pilier fondamental de la mondialisation economique et du tourisme international.")
para("Selon l'Organisation de l'Aviation Civile Internationale (OACI), le nombre de passagers aeriens dans le monde a depasse les 4 milliards en 2019, avant la pandemie de COVID-19. La reprise post-pandemique a demontre la resilience du secteur, avec un retour rapide vers les niveaux de trafic historiques des 2023.")

heading("II.2. La distribution des billets d'avion", 3)
para("La distribution des billets d'avion a traverse plusieurs revolutions technologiques :")
bullet("Premiere generation — Les Systemes de Distribution Globale (GDS) : Apparus dans les annees 1960 et 1970, les GDS tels qu'Amadeus, Sabre et Galileo ont permis aux agences de voyage d'acceder electroniquement aux inventaires des compagnies aeriennes.")
bullet("Deuxieme generation — Le Web : L'avenement d'Internet dans les annees 1990 a permis l'emergence des premieres plateformes de reservation en ligne. Les compagnies aeriennes ont egalement developpe leurs propres sites de vente directe.")
bullet("Troisieme generation — Le Mobile : Depuis le milieu des annees 2010, les applications mobiles dominent l'experience de reservation.")

heading("II.3. Les agences de voyage en ligne (OTA)", 3)
para("Une Online Travel Agency (OTA) est une plateforme numerique permettant aux voyageurs de rechercher, comparer et reserver des services de voyage sans passer par une agence physique. Notre application s'inscrit dans ce modele OTA, jouant le role d'intermediaire entre le voyageur et les compagnies aeriennes via l'API Duffel.")

heading("III. Applications Mobiles", 2)
heading("III.1. Definition", 3)
para("Une application mobile est un logiciel concu pour fonctionner sur un appareil mobile — telephone intelligent ou tablette numerique. Contrairement aux applications de bureau, les applications mobiles exploitent les capacites specifiques des appareils portables : ecran tactile, GPS, appareil photo, capteurs d'orientation, et connectivite permanente.")

heading("III.2. Types d'applications mobiles", 3)
heading("III.2.1. Applications natives", 4)
para("Les applications natives sont developpees specifiquement pour un systeme d'exploitation mobile particulier. Pour Android, on utilise Kotlin ou Java avec Android Studio. Pour iOS, on utilise Swift ou Objective-C avec Xcode. Ces applications offrent les meilleures performances mais necessitent deux codebases distinctes.")

heading("III.2.2. Applications hybrides", 4)
para("Les applications hybrides sont developpees avec des technologies web puis encapsulees dans un conteneur natif. Des frameworks comme Ionic ou Cordova appartiennent a cette categorie. Elles offrent un developpement plus rapide mais des performances inferieures aux applications natives.")

heading("III.2.3. Applications cross-platform", 4)
para("Les frameworks cross-platform permettent d'ecrire une seule codebase qui compile en code natif pour chaque plateforme. Flutter (Google), React Native (Meta) et Xamarin (Microsoft) sont les principaux representants. Flutter se distingue par son moteur de rendu graphique propre garantissant des performances proches des applications natives.")

heading("III.3. Avantages des applications mobiles par rapport aux applications web", 3)
table(
    ["Critere", "Application Mobile", "Application Web"],
    [
        ["Accessibilite", "Disponible hors ligne (partiellement)", "Necessite une connexion constante"],
        ["Performance", "Elevee (acces natif)", "Variable (depend du navigateur)"],
        ["UX", "Optimisee pour le tactile", "Adaptee mais moins fluide"],
        ["Notifications", "Push notifications natives", "Notifications web limitees"],
        ["Acces materiel", "GPS, camera, capteurs", "Limite"],
        ["Fidelisation", "Icone permanente sur l'ecran", "Depend du signet"],
    ]
)

heading("IV. Architecture Client/Serveur", 2)
heading("IV.1. Presentation generale", 3)
para("L'architecture client/serveur est un modele de communication reseau dans lequel certains noeuds (les clients) envoient des requetes a d'autres noeuds (les serveurs) qui y repondent. Dans notre application, cette architecture se decline comme suit : l'application Flutter installée sur le smartphone constitue le client, le backend Django joue le role du serveur, et la base de donnees MySQL assure la persistance des donnees.")

heading("IV.2. API REST", 3)
para("REST (Representational State Transfer) est un style architectural pour la conception de services web, formalise par Roy Fielding dans sa these de doctorat en 2000. Notre API respecte les principes REST : interface uniforme, stateless, cacheable et systeme en couches.")
table(
    ["Methode HTTP", "Usage dans notre API"],
    [
        ["GET",    "Recuperer des ressources (liste de vols, reservations)"],
        ["POST",   "Creer une ressource (nouvelle reservation, connexion)"],
        ["PATCH",  "Modifier partiellement une ressource (statut reservation)"],
        ["DELETE", "Supprimer une ressource (annulation de reservation)"],
    ]
)

heading("IV.3. Authentification par JWT", 3)
para("Notre API utilise le standard JWT (JSON Web Token) pour l'authentification. Le client envoie ses identifiants, le serveur retourne un access token et un refresh token. Le client joint l'access token dans l'en-tete de chaque requete protegee. Quand l'access token expire, le client utilise le refresh token pour en obtenir un nouveau sans se reconnecter.")

heading("IV.4. Integration de l'API Duffel", 3)
para("Duffel est une plateforme technologique britannique qui propose une API moderne permettant aux developpeurs d'acceder directement aux inventaires de vols de plus de 300 compagnies aeriennes mondiales. Notre integration exploite les endpoints suivants :")
table(
    ["Endpoint Duffel", "Usage dans notre application"],
    [
        ["POST /air/offer_requests",       "Recherche de vols avec passagers et dates"],
        ["GET /air/offers/{id}",           "Validation d'une offre avant reservation"],
        ["POST /air/orders",               "Creation d'une commande (emission du billet)"],
        ["POST /air/order_cancellations",  "Demande d'annulation"],
        ["GET /places/suggestions",        "Autocompletion de la recherche d'aeroports"],
    ]
)

heading("V. UML — Langage de Modelisation Unifie", 2)
para("UML (Unified Modeling Language) est un langage de modelisation graphique standardise, normalise par l'OMG. Il permet de representer visuellement la structure et le comportement d'un systeme logiciel. Dans ce memoire, nous utilisons trois types de diagrammes : les diagrammes de cas d'utilisation, les diagrammes de classes et les diagrammes de sequence.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPITRE 2
# ══════════════════════════════════════════════════════════════════════════════
heading("CHAPITRE 2 : ANALYSE ET CONCEPTION", 1)

heading("I. Introduction", 2)
para("Ce chapitre presente l'analyse des besoins de notre application et sa conception detaillee a l'aide des notations UML. Nous commencerons par identifier les acteurs et leurs besoins, puis nous modeliserons le systeme a travers les differents diagrammes UML.")

heading("II. Analyse des Besoins", 2)
heading("II.1. Besoins fonctionnels", 3)

heading("II.1.1. Besoins du Client", 4)
bullet("Authentification : S'inscrire avec un email et un mot de passe ; se connecter et se deconnecter")
bullet("Recherche de vols : Rechercher des vols aller simple ou aller-retour en specif. l'origine, destination, date, passagers et classe")
bullet("Autocompletion : Beneficier d'une suggestion intelligente des aeroports lors de la saisie")
bullet("Reservation : Reserver un vol en fournissant les informations du passager")
bullet("Code promotionnel : Appliquer un code promo pour beneficier d'une reduction")
bullet("Mode de paiement : Choisir entre le paiement ulterieur ou par carte bancaire (CIB/Edahabia)")
bullet("Suivi des reservations : Consulter l'historique et le statut des reservations")
bullet("Billet electronique : Telecharger son billet au format PDF une fois confirme")

heading("II.1.2. Besoins de l'Agent", 4)
bullet("Consulter l'ensemble des reservations de tous les clients")
bullet("Confirmer une reservation en attente, declenchant l'emission du billet via Duffel")
bullet("Annuler une reservation avec remboursement si deja confirmee")
bullet("Effectuer une reservation au nom d'un client")
bullet("Consulter le tableau de bord avec les statistiques globales")

heading("II.1.3. Besoins de l'Administrateur", 4)
bullet("Creer, modifier, bloquer/debloquer et supprimer des comptes utilisateurs")
bullet("Creer et gerer des codes promo (type, valeur, date d'expiration, nombre d'utilisations)")
bullet("Acceder au tableau de bord complet (utilisateurs, reservations, revenus)")

heading("II.2. Besoins non fonctionnels", 3)
bullet("Performance : Les resultats de recherche de vols doivent s'afficher en moins de 10 secondes")
bullet("Securite : Authentification par JWT, validation cote serveur, protection contre les injections SQL")
bullet("Disponibilite : L'application doit etre accessible en continu (24h/24, 7j/7)")
bullet("Compatibilite : L'application mobile doit fonctionner sur Android 8.0 et versions ulterieures")

heading("III. Diagramme de Cas d'Utilisation", 2)
heading("III.1. Identification des acteurs", 3)
para("Notre systeme implique trois acteurs principaux. Le Client est l'utilisateur final qui effectue des recherches et des reservations. L'Agent est l'employe de l'agence qui gere les demandes et peut reserver pour les clients. L'Administrateur est le responsable de la plateforme qui dispose de tous les droits.")

heading("III.2. Relations entre cas d'utilisation", 3)
para("Le diagramme de cas d'utilisation utilise deux types de relations :")
bullet("<<include>> : relation obligatoire — ex: 'Reserver un vol' inclut toujours 'Se connecter'")
bullet("<<extend>> : relation optionnelle — ex: 'Appliquer code promo' etend 'Reserver un vol'")

heading("III.3. Diagramme global", 3)
image(BASE+'uc_0_global.png', "Figure 1 : Diagramme de cas d'utilisation global", width_cm=16)

heading("IV. Diagramme de Classes", 2)
heading("IV.1. Description des classes principales", 3)
table(
    ["Classe", "Attributs principaux", "Role"],
    [
        ["Utilisateur", "id, nom, prenom, email, password, role, telephone", "Classe de base pour tous les profils"],
        ["Reservation", "id, offerId, status, totalAmount, bookingReference", "Represente une reservation de vol"],
        ["Passager", "id, type, prenom, nom, dateNaissance, nationalite", "Passager associe a une reservation"],
        ["Paiement", "id, montant, methode, status", "Paiement lie a une reservation"],
        ["CodePromo", "id, code, typeReduction, valeur, maxUtilisations", "Code de reduction cree par l'admin"],
    ]
)

heading("IV.2. Relations entre classes", 3)
bullet("Utilisateur 1:N Reservation — Un utilisateur peut avoir plusieurs reservations")
bullet("Reservation 1:N Passager — Une reservation contient un ou plusieurs passagers")
bullet("Reservation 1:1 Paiement — Une reservation est associee a un paiement")
bullet("Reservation 0:1 CodePromo — Une reservation peut beneficier d'un code promo")

image(BASE+'diagramme_classes.png', "Figure 2 : Diagramme de classes", width_cm=16)

heading("V. Diagrammes de Sequence", 2)
heading("V.1. Authentification", 3)
para("Ce diagramme decrit le processus de connexion. Le client envoie ses identifiants au systeme, qui interroge la base de donnees. Si l'utilisateur est trouve et le mot de passe correct (alt), les tokens JWT sont generes et retournes. Sinon, une erreur est retournee.")

heading("V.2. Recherche de Vols", 3)
para("Le client envoie les criteres de recherche. Le systeme prepare les criteres et interroge le Service Vols (Duffel). Si des vols sont trouves (alt), une boucle (loop) traite chaque vol pour lui attacher un resume, les resultats sont mis en cache et retournes. Sinon, une liste vide est retournee.")

image(BASE+'seq_1_authentification.png', "Figure 3 : Diagramme de sequence — Authentification", width_cm=14)
image(BASE+'seq_2_inscription.png',      "Figure 4 : Diagramme de sequence — Inscription",      width_cm=14)

heading("V.3. Reservation Client", 3)
para("Apres authentification (ref), le client soumet sa reservation. Si l'offre est disponible (alt), le systeme cree la reservation avec le statut EN_ATTENTE, enregistre les passagers en boucle (loop pour chaque passager), et retourne la confirmation. Si le code promo est valide, la reduction est appliquee.")

image(BASE+'seq_3_recherche_vols.png',   "Figure 5 : Diagramme de sequence — Recherche de vols",    width_cm=15)
image(BASE+'seq_4_reservation_client.png', "Figure 6 : Diagramme de sequence — Reservation client", width_cm=15)

heading("V.4. Confirmation par l'Agent", 3)
para("L'agent consulte les reservations et confirme. Si le statut est EN_ATTENTE (alt), le systeme appelle Duffel pour creer le billet, met a jour la reservation avec le statut CONFIRME et la reference PNR, puis retourne la confirmation a l'agent.")

image(BASE+'seq_5_confirmation_agent.png', "Figure 7 : Diagramme de sequence — Confirmation agent",      width_cm=15)
image(BASE+'seq_6_annulation.png',         "Figure 8 : Diagramme de sequence — Annulation",              width_cm=15)
image(BASE+'seq_7_code_promo.png',         "Figure 9 : Diagramme de sequence — Application code promo",  width_cm=14)
image(BASE+'seq_8_gestion_utilisateurs.png', "Figure 10 : Diagramme de sequence — Gestion utilisateurs", width_cm=14)

heading("VI. Schema de la Base de Donnees", 2)
para("Notre base de donnees MySQL comprend six tables principales :")
image(BASE+'schema_relationnel.png', "Figure 11 : Schema relationnel de la base de donnees", width_cm=16)
table(
    ["Table", "Cle Primaire", "Cles Etrangeres", "Description"],
    [
        ["users",        "id (UUID)", "—",                          "Comptes utilisateurs (admin/agent/client)"],
        ["reservations", "id (UUID)", "user_id → users",            "Reservations de vols"],
        ["passengers",   "id (UUID)", "reservation_id → reservations", "Passagers par reservation"],
        ["payments",     "id (UUID)", "reservation_id → reservations", "Paiements (1:1 avec reservation)"],
        ["promo_codes",  "id (INT)",  "—",                          "Codes promotionnels"],
        ["flight_cache", "offer_id",  "—",                          "Cache temporaire des offres Duffel"],
    ]
)

heading("VI.2. Tables supprimees grace a l'API Duffel", 3)
para("Dans la phase de conception (diagramme de classes), plusieurs entites avaient ete identifiees : Vol, Aeroport, CompagnieAvion, Pays, Ville, Passeport et BilletDAvion. Cependant, lors de l'implementation, ces tables ont ete supprimees de la base de donnees pour la raison suivante.")
para("L'integration de l'API Duffel rend ces tables redondantes et inutiles. En effet, Duffel gere directement ces donnees de son cote : les informations sur les vols, les aeroports, les compagnies aeriennes et les billets sont retournees en temps reel par l'API et stockees sous forme de JSON dans le champ raw_duffel_order de la table reservations.")
table(
    ["Table supprimee", "Raison", "Remplacee par"],
    [
        ["Vol",            "Geree par Duffel en temps reel",     "JSON raw_duffel_order"],
        ["Aeroport",       "Donnees fournies par Duffel",        "Code IATA (VARCHAR)"],
        ["CompagnieAvion", "Incluse dans la reponse Duffel",     "JSON raw_duffel_order"],
        ["BilletDAvion",   "Billet emis et stocke par Duffel",   "booking_reference + JSON"],
        ["Passeport",      "Champs integres dans Passenger",     "Colonnes de passengers"],
        ["Pays / Ville",   "Non necessaires (code IATA suffisant)", "VARCHAR iata_code"],
    ]
)
para("Cette approche presente plusieurs avantages : elle evite la duplication de donnees entre notre base et les serveurs Duffel, reduit la complexite de maintenance, et garantit que les informations affichees sont toujours synchronisees avec les donnees reelles des compagnies aeriennes. C'est une pratique courante dans les systemes qui s'appuient sur des API tierces specialisees.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPITRE 3
# ══════════════════════════════════════════════════════════════════════════════
heading("CHAPITRE 3 : REALISATION", 1)

heading("I. Introduction", 2)
para("Ce chapitre presente la realisation concrete de notre application. Nous decrirons les environnements de developpement, l'architecture implementee, et le fonctionnement des differents modules avec leurs flux de traitement.")

heading("II. Environnement de Developpement", 2)
heading("II.1. Environnement materiel", 3)
bullet("Processeur : Intel Core i5 / AMD Ryzen 5")
bullet("Memoire RAM : 8 Go minimum")
bullet("Systeme d'exploitation : Windows 11 Pro")
bullet("Appareil de test : Smartphone Android pour les tests sur device physique")

heading("II.2. Flutter", 3)
para("Flutter est un framework open-source developpe par Google permettant de creer des applications natives compilees pour mobile, web et bureau a partir d'une seule base de code. Il utilise le langage Dart compile en code natif ARM et dispose d'un moteur de rendu graphique propre (Skia/Impeller) independant des widgets natifs de chaque plateforme.")
table(
    ["Package Flutter", "Version", "Usage"],
    [
        ["http",               "^1.2.2",  "Requetes HTTP vers le backend Django"],
        ["shared_preferences", "^2.2.3",  "Stockage local des tokens JWT"],
        ["pdf",                "^3.11.0", "Generation des billets electroniques PDF"],
        ["printing",           "^5.13.1", "Telechargement et impression des PDF"],
    ]
)

heading("II.3. Django REST Framework", 3)
para("Django est un framework web Python suivant le patron MVT (Model-View-Template). Django REST Framework (DRF) est une extension facilitant la creation d'API RESTful. Nous utilisons notamment les Serializers pour la validation et conversion des donnees, les APIView pour definir les endpoints, le systeme de Permissions pour le controle d'acces, et SimpleJWT pour l'authentification par tokens.")

heading("II.4. MySQL", 3)
para("MySQL est un systeme de gestion de base de donnees relationnelle open-source. Il garantit la coherence des donnees via les contraintes d'integrite referentielle, supporte les transactions ACID pour la fiabilite des operations de reservation, et offre un support natif du type JSON pour stocker les donnees brutes Duffel.")

heading("II.5. API Duffel", 3)
para("Duffel est une API de distribution aerienne de nouvelle generation permettant d'acceder aux inventaires de plus de 300 compagnies aeriennes via une interface REST moderne. Elle offre un environnement sandbox pour les tests et une documentation claire, ce qui la rend ideale pour un projet academique.")

heading("II.6. Outils de developpement", 3)
table(
    ["Outil", "Usage"],
    [
        ["Visual Studio Code", "Editeur de code principal avec extensions Flutter/Dart"],
        ["Android Studio",     "Emulateur Android et debogage mobile"],
        ["MySQL Workbench",    "Administration et visualisation de la base de donnees"],
        ["Postman",            "Test et documentation des endpoints API"],
        ["Git / GitHub",       "Controle de version et hebergement du code source"],
    ]
)

heading("III. Architecture de l'Application", 2)
heading("III.1. Architecture trois tiers", 3)
para("Notre application suit une architecture trois tiers. Le premier tier est la couche Presentation : l'application Flutter avec ses ecrans mobiles, sa navigation et ses composants UI. Le deuxieme tier est la couche Logique Metier : le backend Django avec ses vues, serializers, services et le client Duffel. Le troisieme tier est la couche Donnees : la base de donnees MySQL.")

heading("III.2. Structure du projet Flutter", 3)
bullet("lib/main.dart : Point d'entree, AuthWrapper, routing selon le role")
bullet("lib/api.dart : Client API centralisant toutes les requetes HTTP")
bullet("lib/models/ : Modeles de donnees (Flight, Booking)")
bullet("lib/navigation/ : Navigations par role (MainNavigation, AdminNavigation, AgentNavigation)")
bullet("lib/screens/ : Ecrans organises par domaine (auth, home, flights, trips, admin, agent)")
bullet("lib/utils/ : Utilitaires (generation PDF des billets)")
bullet("lib/widgets/ : Composants reutilisables (recherche aeroport)")
bullet("assets/images/ : Ressources statiques (carte du monde)")

heading("III.3. Structure du projet Django", 3)
bullet("apps/users/ : Modele User personnalise, authentification, gestion des comptes")
bullet("apps/reservations/ : Reservations, Passagers, Paiements, Codes promo")
bullet("apps/flights/ : Recherche de vols, recherche d'aeroports")
bullet("core/duffel_client.py : Client HTTP vers l'API Duffel")
bullet("core/permissions.py : Classes de permissions personnalisees")
bullet("config/settings.py : Configuration Django")
bullet("config/urls.py : Routage URL principal")

heading("IV. Fonctionnement des Modules", 2)
heading("IV.1. Authentification et Gestion des Roles", 3)
para("Notre systeme implemente une authentification basee sur JWT avec trois roles distincts. A l'initialisation de l'application, le composant AuthWrapper verifie si un access token est stocke localement. Si oui, il recupere le profil pour determiner le role et dirige l'utilisateur vers l'interface correspondante : AdminNavigation pour admin (5 onglets), AgentNavigation pour agent (4 onglets), ou MainNavigation pour client (3 onglets).")

heading("IV.2. Recherche de Vols", 3)
para("L'utilisateur saisit l'origine et la destination via le widget d'autocompletion (connecte a l'API Duffel en temps reel), selectionne les dates et configure les passagers avec leurs types et ages. Le backend construit le tableau de passagers Duffel avec les ages des enfants, appelle l'API via POST /air/offer_requests, enrichit chaque offre avec un resume calcule, et retourne les resultats au client.")
para("Pour les vols aller-retour, deux slices sont envoyees a Duffel : la premiere pour le vol aller et la seconde pour le vol retour, permettant d'obtenir des offres combinees.")

heading("IV.3. Flux de Reservation en Deux Etapes", 3)
para("Notre application implemente un flux de reservation en deux etapes adapte au modele des agences de voyage algeriennes.")
para("Etape 1 — Reservation client : Le client remplit le formulaire passager, applique eventuellement un code promo, et choisit le mode de paiement. Le backend valide les donnees, verifie l'offre aupres de Duffel, et cree la reservation avec le statut PENDING. Les donnees sont stockees dans pending_booking_data en attendant la confirmation.", bold=False)
para("Etape 2 — Confirmation agent : L'agent examine la demande et clique sur Confirmer. Le backend recupere les donnees stockees et appelle Duffel POST /air/orders. Duffel emet le billet et retourne le numero PNR. La reservation passe au statut CONFIRMED.", bold=False)
para("Exception — Paiement par carte : Lorsque le client choisit le paiement par carte (CIB/Edahabia), le backend appelle immediatement Duffel et confirme la reservation sans intervention humaine.")

heading("IV.4. Generation du Billet PDF", 3)
para("Une fois la reservation confirmee, le client peut telecharger son billet electronique en PDF. Ce document comprend la reference PNR mise en evidence, le trajet avec les codes IATA, la liste des passagers, le montant total paye et les mentions legales. La generation est realisee cote Flutter grace au package pdf, sans appel supplementaire au backend.")

heading("IV.5. Codes Promotionnels", 3)
para("L'administrateur cree des codes promo caracterises par un code unique, un type de reduction (pourcentage ou montant fixe), une valeur, un nombre maximum d'utilisations et une date d'expiration. Lors de l'application, le backend verifie la validite du code et calcule le montant reduit. Apres confirmation, le compteur d'utilisation est incremente automatiquement.")

heading("V. Algorithme A priori — Systeme de Recommandation", 2)
heading("V.1. Presentation de l'algorithme A priori", 3)
para("L'algorithme A priori est un algorithme de fouille de donnees (data mining) propose par Agrawal et Srikant en 1994. Il permet d'extraire des regles d'association a partir d'un ensemble de transactions, c'est-a-dire de decouvrir des relations entre des elements qui apparaissent frequemment ensemble dans les donnees.")
para("Dans notre application, nous appliquons cet algorithme sur l'historique des reservations confirmees afin de generer des recommandations intelligentes de destinations pour les utilisateurs. Ce systeme de recommandation personnalise l'experience de recherche en suggerant automatiquement les destinations les plus pertinentes en fonction des habitudes de voyage de l'ensemble des clients.")

heading("V.2. Concepts fondamentaux", 3)
para("L'algorithme repose sur trois metriques essentielles :")
bullet("Support : frequence d'apparition d'un ensemble de destinations dans toutes les transactions. Un support de 20% signifie que 20% des clients ont reserve cette destination.")
bullet("Confiance : probabilite conditionnelle que la destination B soit reservee sachant que la destination A l'a ete. Une confiance de 80% indique que 80% des clients ayant reserve A ont aussi reserve B.")
bullet("Lift : mesure la force reelle de la regle par rapport au hasard. Un lift superieur a 1 indique une correlation positive et significative entre les destinations.")
table(
    ["Metrique", "Formule", "Interpretation"],
    [
        ["Support",    "nb(A et B) / nb total",  "Frequence de la paire dans les donnees"],
        ["Confiance",  "Support(A et B) / Support(A)", "Probabilite de B si A est present"],
        ["Lift",       "Confiance(A->B) / Support(B)", "Force de la regle vs le hasard"],
    ]
)

heading("V.3. Application dans notre systeme", 3)
para("Le flux d'execution de l'algorithme dans notre application se deroule en cinq etapes :")
bullet("Etape 1 — Extraction des donnees : Le systeme recupere toutes les reservations confirmees depuis la base de donnees MySQL, en extrayant les codes IATA d'origine et de destination de chaque reservation.")
bullet("Etape 2 — Construction des transactions : Les reservations sont regroupees par utilisateur pour former des transactions. Chaque transaction represente l'ensemble des trajets effectues par un client.")
bullet("Etape 3 — Encodage et application A priori : Les transactions sont encodees au format binaire (TransactionEncoder), puis l'algorithme A priori est applique avec un seuil de support minimum de 5% et de confiance minimum de 30%.")
bullet("Etape 4 — Filtrage des regles : Seules les regles dont l'antecedent correspond a l'origine saisie par l'utilisateur sont conservees et triees par confiance decroissante.")
bullet("Etape 5 — Affichage dans Flutter : Les cinq meilleures destinations recommandees sont affichees sous le formulaire de recherche, avec leur taux de confiance.")
para("Ce module est accessible via l'endpoint GET /api/v1/recommendations/?origin=ALG et retourne une liste de destinations avec leurs metriques de support, confiance et lift. L'interface Flutter permet a l'utilisateur de selectionner directement une destination recommandee pour pre-remplir le formulaire de recherche.")

heading("V.4. Exemple de resultat", 3)
table(
    ["Origine", "Destination recommandee", "Support", "Confiance", "Lift"],
    [
        ["ALG", "CDG (Paris)",     "35%", "82%", "2.1"],
        ["ALG", "TUN (Tunis)",     "28%", "75%", "1.9"],
        ["ALG", "IST (Istanbul)",  "18%", "60%", "1.5"],
    ]
)
para("Interpretation : 82% des clients ayant reserve un vol ALG-CDG ont egalement reserve un vol CDG-ALG (retour), ce qui suggere automatiquement le vol retour comme destination.")

heading("VI. Tests et Validation", 2)
table(
    ["Fonctionnalite", "Resultat"],
    [
        ["Inscription / Connexion",             "Fonctionnel"],
        ["Recherche aller simple",              "Fonctionnel"],
        ["Recherche aller-retour",              "Fonctionnel"],
        ["Autocompletion aeroports (Duffel)",   "Fonctionnel"],
        ["Reservation avec adulte",             "Fonctionnel"],
        ["Reservation avec enfant (age requis)","Fonctionnel"],
        ["Application code promo",              "Fonctionnel"],
        ["Paiement ulterieur (PENDING)",        "Fonctionnel"],
        ["Paiement CIB (confirmation auto)",    "Fonctionnel"],
        ["Confirmation agent",                  "Fonctionnel"],
        ["Annulation reservation",              "Fonctionnel"],
        ["Telechargement billet PDF",           "Fonctionnel"],
        ["Gestion utilisateurs (admin)",        "Fonctionnel"],
        ["Gestion codes promo (admin)",         "Fonctionnel"],
        ["Recommandations A priori",            "Fonctionnel"],
        ["Telechargement billet PDF",           "Fonctionnel"],
    ]
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
heading("CONCLUSION GENERALE", 1)
para("Ce projet de fin d'etudes nous a permis de concevoir et developper une application mobile complete de reservation de vols, mettant en oeuvre un ensemble coherent de technologies modernes du developpement logiciel.")
para("Sur le plan technique, nous avons reussi a integrer harmonieusement Flutter pour l'interface mobile cross-platform, Django REST Framework pour le backend API, MySQL pour la persistance des donnees, et l'API Duffel pour l'acces en temps reel aux inventaires de vols des compagnies aeriennes mondiales.")
para("Sur le plan fonctionnel, notre application repond aux besoins identifies lors de l'analyse : les clients peuvent rechercher et reserver des vols avec une experience fluide et intuitive, les agents disposent d'une interface dediee pour confirmer les billets, et les administrateurs beneficient d'outils complets pour superviser la plateforme.")
para("Parmi les choix architecturaux notables, le flux de reservation en deux etapes (PENDING -> CONFIRMED) s'est revele particulierement adapte au contexte des agences de voyage algeriennes. L'integration du paiement par carte bancaire (CIB/Edahabia) avec confirmation automatique offre egalement une alternative moderne.")
para("Ce projet nous a permis d'acquerir des competences precieuses en developpement mobile avec Flutter, conception d'API RESTful avec Django, modelisation UML, et integration avec des API tierces professionnelles.")

heading("Perspectives d'amelioration", 2)
bullet("Notifications push : Alerter le client lors de la confirmation de sa reservation")
bullet("Paiement en ligne integre : Integration complete avec une passerelle CIB/Edahabia")
bullet("Gestion des bagages : Permettre la selection de bagages supplementaires")
bullet("Multi-langues : Support de l'arabe et de l'anglais")
bullet("Deploiement cloud : Hebergement sur infrastructure cloud pour une disponibilite professionnelle")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
heading("REFERENCES BIBLIOGRAPHIQUES", 1)
refs = [
    # Ouvrages
    "[1] RICHARDSON, Leonard. RESTful Web APIs. O'Reilly Media, 2013. ISBN 978-1-449-35820-9.",
    "[2] GREENFELD, Daniel ; ROY, Audrey. Two Scoops of Django 3.x : Best Practices for the Django Web Framework. Two Scoops Press, 2021.",
    "[3] WINDMILL, Eric. Flutter in Action. Manning Publications, 2020. ISBN 978-1-617-29647-9.",
    "[4] FOWLER, Martin. Patterns of Enterprise Application Architecture. Addison-Wesley, 2002. ISBN 978-0-321-12742-6.",
    "[5] AGRAWAL, Rakesh ; SRIKANT, Ramakrishnan. Fast Algorithms for Mining Association Rules. Proceedings of the 20th VLDB Conference, Santiago, Chile, 1994.",
    "[6] HAN, Jiawei ; KAMBER, Micheline ; PEI, Jian. Data Mining : Concepts and Techniques. 3e edition. Morgan Kaufmann, 2011. ISBN 978-0-123-81479-1.",
    # Documentation technique
    "[7] Google LLC. Flutter Documentation officielle. Disponible sur : https://docs.flutter.dev. Consulte en 2024.",
    "[8] Django Software Foundation. Django Documentation v4.x. Disponible sur : https://docs.djangoproject.com. Consulte en 2024.",
    "[9] Encode Ltd. Django REST Framework Documentation. Disponible sur : https://www.django-rest-framework.org. Consulte en 2024.",
    "[10] Duffel Technologies Ltd. Duffel API v2 Documentation. Disponible sur : https://duffel.com/docs. Consulte en 2024.",
    "[11] MySQL AB / Oracle Corporation. MySQL 8.0 Reference Manual. Disponible sur : https://dev.mysql.com/doc. Consulte en 2024.",
    "[12] Rasbt. MLxtend : Machine Learning Extensions for Python. Disponible sur : https://rasbt.github.io/mlxtend. Consulte en 2024.",
    # Articles et normes
    "[13] IETF. RFC 7519 — JSON Web Token (JWT). Disponible sur : https://tools.ietf.org/html/rfc7519. Mai 2015.",
    "[14] FIELDING, Roy Thomas. Architectural Styles and the Design of Network-based Software Architectures. These de doctorat, Universite de Californie, Irvine, 2000.",
    "[15] Object Management Group (OMG). Unified Modeling Language (UML) Specification Version 2.5.1. Disponible sur : https://www.omg.org/spec/UML. 2017.",
    "[16] OACI — Organisation de l'Aviation Civile Internationale. Annual Report of the ICAO Council 2023. Montreal, Canada, 2023.",
    "[17] DART TEAM — Google. Dart Programming Language. Disponible sur : https://dart.dev. Consulte en 2024.",
]
for ref in refs:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    set_font(run, size=11)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0)

# ── numerotation des pages ────────────────────────────────────────────────────
add_page_numbers(doc)

# ── sauvegarde ────────────────────────────────────────────────────────────────
path = 'c:/Users/tassili/Desktop/pfe/Memoire_PFE_Reservation_Vols.docx'
doc.save(path)
print(f"Sauvegarde: {path}")
